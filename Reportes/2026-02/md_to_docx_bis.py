import markdown
from docx import Document
import os

md_path = r"c:/VLP/GitHub/ACTIVIDADES_2026/Reportes/2026-02/Informe_Mensual_2026-02-bis.md"
docx_path = r"c:/VLP/GitHub/ACTIVIDADES_2026/Reportes/2026-02/Informe_Mensual_2026-02-bis.docx"

def md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()
    doc = Document()
    doc.add_heading('Informe Mensual de Actividades (2026-02-bis)', 0)
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. '):
            doc.add_paragraph(line, style='List Number')
        elif line.strip().startswith('|') and line.strip().endswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            if len(table_lines) >= 2:
                headers = [h.strip() for h in table_lines[0].strip('|').split('|')]
                rows = [
                    [cell.strip() for cell in row.strip('|').split('|')]
                    for row in table_lines[2:]
                ]
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for j, h in enumerate(headers):
                    hdr_cells[j].text = h
                for row in rows:
                    row_cells = table.add_row().cells
                    for j, cell in enumerate(row):
                        row_cells[j].text = cell
            continue
        elif line.strip() == '':
            doc.add_paragraph('')
        else:
            doc.add_paragraph(line)
        i += 1
    doc.save(docx_path)
    print(f"Documento Word generado en: {docx_path}")

if __name__ == "__main__":
    md_to_docx(md_path, docx_path)
