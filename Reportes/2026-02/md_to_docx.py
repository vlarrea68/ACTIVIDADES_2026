
import re
from docx import Document

md_path = r"Reportes/2026-02/Informe_Mensual_2026-02.md"
docx_path = r"Reportes/2026-02/Informe_Mensual_2026-02.docx"

def parse_table(lines, start):
    table_lines = []
    i = start
    while i < len(lines) and lines[i].strip().startswith('|'):
        table_lines.append(lines[i].strip())
        i += 1
    if len(table_lines) < 2:
        return None, i
    headers = [h.strip() for h in table_lines[0].strip('|').split('|')]
    rows = [
        [cell.strip() for cell in row.strip('|').split('|')]
        for row in table_lines[2:]
    ]
    return (headers, rows), i

def md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()
    doc = Document()
    doc.add_heading('Informe Mensual de Actividades (2026-02)', 0)
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. '):
            doc.add_paragraph(line, style='List Number')
        elif line.strip().startswith('|'):
            # Detect and parse markdown table
            table, next_i = parse_table(lines, i)
            if table:
                headers, rows = table
                table_obj = doc.add_table(rows=1, cols=len(headers))
                table_obj.style = 'Table Grid'
                hdr_cells = table_obj.rows[0].cells
                for j, h in enumerate(headers):
                    hdr_cells[j].text = h
                for row in rows:
                    row_cells = table_obj.add_row().cells
                    for j, cell in enumerate(row):
                        row_cells[j].text = cell
                i = next_i
                continue
        elif line.strip() == '':
            doc.add_paragraph('')
        else:
            doc.add_paragraph(line)
        i += 1
        doc.save(docx_path)
    print(f"Documento Word generado en: {docx_path}")

if __name__ == "__main__":
    md_to_docx(md_path, docx_path)
