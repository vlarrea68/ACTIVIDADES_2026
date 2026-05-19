from pathlib import Path

from docx import Document


BASE_DIR = Path(__file__).resolve().parent
MD_PATH = BASE_DIR / "Informe_Mensual_2026-06.md"
DOCX_PATH = BASE_DIR / "Informe_Mensual_2026-06.docx"


def parse_table(lines, start_index):
    table_lines = []
    index = start_index
    while index < len(lines) and lines[index].strip().startswith("|"):
        table_lines.append(lines[index].strip())
        index += 1

    if len(table_lines) < 2:
        return None, index

    headers = [cell.strip() for cell in table_lines[0].strip("|").split("|")]
    rows = [
        [cell.strip() for cell in row.strip("|").split("|")]
        for row in table_lines[2:]
    ]
    return (headers, rows), index


def markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    markdown_text = md_path.read_text(encoding="utf-8")
    document = Document()
    lines = markdown_text.splitlines()
    index = 0

    while index < len(lines):
        line = lines[index]

        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("1. "):
            document.add_paragraph(line, style="List Number")
        elif line.strip().startswith("|"):
            table, next_index = parse_table(lines, index)
            if table:
                headers, rows = table
                table_obj = document.add_table(rows=1, cols=len(headers))
                table_obj.style = "Table Grid"
                for column, header in enumerate(headers):
                    table_obj.rows[0].cells[column].text = header
                for row in rows:
                    cells = table_obj.add_row().cells
                    for column, value in enumerate(row):
                        cells[column].text = value
                index = next_index
                continue
        elif line.strip().startswith("```"):
            index += 1
            code_lines = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            document.add_paragraph("\n".join(code_lines))
        elif line.strip() == "---":
            document.add_paragraph("-" * 40)
        elif line.strip() == "":
            document.add_paragraph("")
        else:
            document.add_paragraph(line)

        index += 1

    document.save(docx_path)
    print(f"Documento Word generado en: {docx_path}")


if __name__ == "__main__":
    markdown_to_docx(MD_PATH, DOCX_PATH)